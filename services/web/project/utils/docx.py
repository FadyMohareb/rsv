"""
docx.py
=======
This utilities module is in charge of processing the data of a distribution and preparing its docx report.

Functions:

    generate_two_plots(sample_name, sample_data, role)
        Generates two vertical bar plots from sample data: a percentage plot (stacked bars for Genome Coverage and Ns, and offset bars for Similarity) and a read coverage plot (black bars).

    generate_aggregated_plot_by_platform(sample_name, sample_data)
        Generates two vertical bar plots with aggregated (average) metrics by sequencing platform.

    create_pygenometracks_plot(reference_genome, annotation, region, bed_path, bigwig_file, bigwig_consensus_file, output_dir, sample_name, user_lab)
        Uses pyGenomeTracks to generate genome coverage plots.

    create_element(name)
        Creates an XML element for DOCX formatting.

    create_attribute(element, name, value)
        Adds an attribute to an XML element.

    add_page_number(run)
        Inserts a page number field in a DOCX document.
    
    generate_docx_report(report_data, base_dir, role, user_lab, distribution)
        Generates a DOCX report summarizing viric genome analysis results.

:author: Kevin
:version: 0.0.1
:date: 2025-02-21
"""

import datetime
from io import BytesIO
import docx
from docx import Document
from docx.shared import Inches,Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from project.utils.sql_models import Distribution, Organization, Submission
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, ns, parse_xml
from docx.oxml.ns import nsdecls, qn
from project.utils.report_parser import process_all_reports
import matplotlib.pyplot as plt
import numpy as np
import os
import subprocess
import shutil
import copy

def generate_two_plots(sample_name, sample_data, role):
    """
    Generates two vertical bar plots from sample data:
      1. A percentage plot (stacked bars for Genome Coverage and Ns, and offset bars for Similarity).
      2. A read coverage plot (black bars).
    
    :param sample_name: The name of the sample, used for file naming.
    :type sample_name: str
    :param sample_data: Dictionary containing sample statistics for each lab.
                        Each value is expected to be a dict with keys:
                        "coverage", "similarity", "Ns", "Mean coverage depth", and optionally "sequencing_platform".
    :type sample_data: dict
    :param role: User role (affects visualization style, if needed).
    :type role: str
    :return: A tuple with the file paths of the generated percentage plot and read coverage plot.
    :rtype: (str, str)
    """
    # Extract the lab names and anonymize them (for display)
    labs = list(sample_data.keys())
    anonymized_labs = [lab if lab != 'Reference' else 'Reference Lab (NEQAS)' for lab in labs]
    
    # Extract percentage metrics
    genome_coverage_percent = [
        sample_data[lab]["coverage"] * 100 if sample_data[lab]["coverage"] != "N/A" else 0 
        for lab in labs
    ]
    similarity_percent = [
        sample_data[lab]["similarity"] if sample_data[lab]["coverage"] != "N/A" else 0 
        for lab in labs
    ]
    num_ns_percent = [
        sample_data[lab]["Ns"] if sample_data[lab]["coverage"] != "N/A" else 0 
        for lab in labs
    ]
    
    # Extract read coverage metric
    read_coverage = [
        sample_data[lab]["Mean coverage depth"] if sample_data[lab]["Mean coverage depth"] != "N/A" else 0 
        for lab in labs
    ]
    
    # Plot 1: Percentage Metrics Plot
    fig1, ax1 = plt.subplots(figsize=(16, 8))
    bar_width = 0.3
    x_pos = np.arange(len(labs))
    
    # Stacked bars for Genome Coverage and Ns
    ax1.bar(x_pos, genome_coverage_percent, bar_width, 
            label="Genome Coverage (%)", color="#1E3A5F", zorder=3)
    ax1.bar(x_pos, num_ns_percent, bar_width, 
            bottom=genome_coverage_percent, label="Ns in Sequence (%)", color="#FBC02D", zorder=3)
    
    # Offset bars for Similarity
    ax1.bar(x_pos + bar_width, similarity_percent, bar_width, 
            label="Similarity (%)", color="#F57C00", zorder=3)
    
    ax1.set_xlabel("Participant", fontsize=18)
    ax1.set_ylabel("Percentage (%)", fontsize=18)
    ax1.set_xticks(x_pos + bar_width / 2)
    ax1.set_xticklabels(anonymized_labs, rotation=45, ha="right", fontsize=16)
    
    # Horizontal threshold lines, if desired
    ax1.axhline(90, color="grey", linestyle="--", linewidth=1, zorder=-1)
    ax1.axhline(95, color="blue", linestyle="--", linewidth=1, zorder=-1)
    
    # Combine legend from ax1 (percentage plot)
    handles1, labels1 = ax1.get_legend_handles_labels()
    ax1.legend(handles=handles1, labels=labels1, loc="upper center", 
               bbox_to_anchor=(0.5, 1.2), ncol=2, fontsize=14, bbox_transform=ax1.transAxes)
    
    fig1.tight_layout()
    percentage_plot_file = f"{sample_name}_percentage_plot.png"
    plt.savefig(percentage_plot_file, dpi=300, bbox_inches="tight")
    plt.close(fig1)
    
    # Plot 2: Read Coverage Plot (Black Bars)
    fig2, ax2 = plt.subplots(figsize=(16, 8))
    bar_width2 = 0.5
    x_pos2 = np.arange(len(labs))
    
    # Create black bars for read coverage
    ax2.bar(x_pos2, read_coverage, bar_width2, color="grey", label="Read Coverage (Mean)")
    ax2.axhline(50, color="grey", linestyle="--", linewidth=1, zorder=-1)
    ax2.set_xlabel("Participant", fontsize=18)
    ax2.set_ylabel("Read Coverage (Mean)", fontsize=18)
    ax2.set_xticks(x_pos2)
    ax2.set_xticklabels(anonymized_labs, rotation=45, ha="right", fontsize=16)
    ax2.legend(loc="upper center", bbox_to_anchor=(0.5, 1.2), fontsize=14)
    
    fig2.tight_layout()
    read_coverage_plot_file = f"{sample_name}_read_coverage_plot.png"
    plt.savefig(read_coverage_plot_file, dpi=300, bbox_inches="tight")
    plt.close(fig2)
    
    return percentage_plot_file, read_coverage_plot_file

def generate_aggregated_plot_by_platform(sample_name, sample_data, user_lab, role):
    """
    Generates two vertical bar plots with aggregated (average) metrics by sequencing platform.
    The first plot shows percentage metrics (stacked bars for Genome Coverage and Ns, with offset bars for Similarity).
    The second plot shows Read Coverage as black bars.
    The sequencing platform used by user_lab is highlighted by placing it first in the order.
    
    :param sample_name: The name of the sample, used for file naming.
    :type sample_name: str
    :param sample_data: Dictionary containing sample statistics for each lab.
                        Each value is expected to be a dict with keys:
                        "coverage", "similarity", "Ns", "Mean coverage depth", and "sequencing_platform".
    :type sample_data: dict
    :param user_lab: The lab identifier to highlight in the plot.
    :type user_lab: str
    :return: A tuple with the file paths of the generated percentage plot, read coverage plot, and the user platform.
    :rtype: (str, str, str)
    """
    # Group data by sequencing platform
    grouped = {}
    annotation_label="Your Platform" if role=="user" else "Reference lab platform"
    user_platforms = []  # Store user platforms (handles multiple entries)
    for lab, metrics in sample_data.items():
        platforms = metrics.get("sequencing_platform", "Unknown").split(",")  # Split into multiple platforms
        platforms = [p.strip() for p in platforms]  # Remove whitespace

        if lab == user_lab:
            user_platforms=platforms  # Store user lab platforms

        for platform in platforms:
            if platform not in grouped:
                grouped[platform] = {
                    "coverage": [],
                    "similarity": [],
                    "Ns": [],
                    "read_cov": []
                }
            grouped[platform]["coverage"].append(metrics["coverage"] * 100 if metrics["coverage"] != "N/A" else 0)
            grouped[platform]["similarity"].append(metrics["similarity"] if metrics["similarity"] != "N/A" else 0)
            grouped[platform]["Ns"].append(metrics["Ns"] if metrics["Ns"] != "N/A" else 0)
            grouped[platform]["read_cov"].append(metrics["Mean coverage depth"] if metrics["Mean coverage depth"] != "N/A" else 0)
    
    # Reorder platforms: user's platform first (if found), then alphabetical.
    all_platforms = list(grouped.keys())
    matching_platforms = sorted([p for p in user_platforms if p in all_platforms])  # User platforms first
    remaining_platforms = sorted([p for p in all_platforms if p not in user_platforms])  # Others alphabetically
    sorted_platforms = matching_platforms + remaining_platforms  # Final order
    
    # Build lists for plotting and include the submission counts in labels.
    platforms = []
    platformNames = []
    avg_coverage = []
    avg_similarity = []
    avg_ns = []
    avg_read_cov = []
    
    for platform in sorted_platforms:
        vals = grouped[platform]
        count = len(vals["coverage"])
        platforms.append(f"{platform} ({count})")
        platformNames.append(platform)
        avg_coverage.append(np.mean(vals["coverage"]))
        avg_similarity.append(np.mean(vals["similarity"]))
        avg_ns.append(np.mean(vals["Ns"]))
        avg_read_cov.append(np.mean(vals["read_cov"]))
    
    # --------------------------
    # Plot 1: Percentage Metrics Plot
    # --------------------------
    fig1, ax1 = plt.subplots(figsize=(16, 8))
    bar_width = 0.3
    x_pos = np.arange(len(platforms))
    
    # Stacked bars: Genome Coverage and Ns
    ax1.bar(x_pos, avg_coverage, bar_width, label="Genome Coverage (%)", color="#1E3A5F", zorder=3)
    ax1.bar(x_pos, avg_ns, bar_width, bottom=avg_coverage, label="Ns in Sequence (%)", color="#FBC02D", zorder=3)
    # Offset bars for Similarity
    ax1.bar(x_pos + bar_width, avg_similarity, bar_width, label="Similarity (%)", color="#F57C00", zorder=3)
    
    ax1.set_xlabel("Sequencing Platform", fontsize=18)
    ax1.set_ylabel("Percentage (%)", fontsize=18)
    ax1.set_xticks(x_pos + bar_width / 2)
    ax1.set_xticklabels(platforms, rotation=45, ha="right", fontsize=16)
    
    # Optional horizontal thresholds
    ax1.axhline(90, color="grey", linestyle="--", linewidth=1, zorder=-1)
    ax1.axhline(95, color="blue", linestyle="--", linewidth=1, zorder=-1)

    # Highlight the user_lab's sequencing platform
    for i in user_platforms:
        if i in platformNames:
            index = platformNames.index(i)
            ax1.annotate(
                annotation_label, 
                xy=(x_pos[index], max(avg_coverage[index] + avg_ns[index], avg_similarity[index]) + 5),
                xytext=(x_pos[index], max(avg_coverage[index] + avg_ns[index], avg_similarity[index]) + 15),
                arrowprops=dict(facecolor='red', arrowstyle="->"),
                ha="center",
                fontsize=14,
                color="red"
            )
    
    # Legend for percentage plot
    handles1, labels1 = ax1.get_legend_handles_labels()
    ax1.legend(handles=handles1, labels=labels1, loc="upper center", 
               bbox_to_anchor=(0.5, 1.2), ncol=2, fontsize=14, bbox_transform=ax1.transAxes)
    
    fig1.tight_layout()
    percentage_plot_file = f"{sample_name}_aggregated_percentage_plot.png"
    plt.savefig(percentage_plot_file, dpi=300, bbox_inches="tight")
    plt.close(fig1)
    
    # --------------------------
    # Plot 2: Read Coverage Plot
    # --------------------------
    fig2, ax2 = plt.subplots(figsize=(16, 8))
    bar_width2 = 0.5
    x_pos2 = np.arange(len(platforms))
    
    ax2.bar(x_pos2, avg_read_cov, bar_width2, color="grey", label="Read Coverage (Mean)")
    ax2.set_yscale("log")
    ax2.axhline(50, color="red", linestyle="--", linewidth=1, zorder=-1)
    ax2.set_xlabel("Sequencing Platform", fontsize=18)
    ax2.set_ylabel("Read Coverage (Mean)", fontsize=18)
    ax2.set_xticks(x_pos2)  # Explicitly set ticks before labels
    ax2.set_xticklabels(platforms, rotation=45, ha="right", fontsize=16)
    # Dynamically determine meaningful y-ticks
    y_max = np.nanmax(avg_read_cov)
    y_ticks = np.geomspace(50, y_max, num=6)  # Create log-spaced ticks
    y_ticks = np.round(y_ticks, decimals=2)  # Round for better readability

    # Set y-ticks using actual data values
    ax2.set_yticks(y_ticks)
    ax2.set_yticklabels([str(int(t)) if t >= 1 else f"{t:.2f}" for t in y_ticks], fontsize=14)

    ax2.legend(loc="upper center", bbox_to_anchor=(0.5, 1.2), fontsize=14)

    # Highlight the user_lab's sequencing platform
    for user_platform in user_platforms:
        if user_platform in platformNames:
            index = platformNames.index(user_platform)
            y_position = avg_read_cov[index] * 1.2 if avg_read_cov[index]>10 else 10 # Move slightly above the bar
            y_text_position = avg_read_cov[index] * 2.5  if avg_read_cov[index]>10 else 20# Text even higher
            ax2.annotate(
                annotation_label,
                xy=(x_pos2[index], y_position),  
                xytext=(x_pos2[index], y_text_position),  
                arrowprops=dict(facecolor='red', arrowstyle="->"),
                ha="center",
                fontsize=14,
                color="red"
            )
            print(x_pos2[index], y_position)
    
    fig2.tight_layout()
    read_coverage_plot_file = f"{sample_name}_aggregated_readcov_plot.png"
    plt.savefig(read_coverage_plot_file, dpi=300, bbox_inches="tight")
    plt.close(fig2)
    
    return percentage_plot_file, read_coverage_plot_file, user_platforms

def create_pygenometracks_plot(reference_genome, annotation, region, bed_path, bigwig_file, bigwig_consensus_file, output_dir, sample_name, user_lab):
    """
    Generates a genome coverage plot using pyGenomeTracks for a specific user_lab.

    Creates a tracks.ini file that configures:
      - A reference track (FASTA),
      - An annotation track (GFF3),
      - A bigwig track (BigWig file).

    :param reference_genome: Path to the reference genome FASTA file.
    :type reference_genome: str
    :param annotation: Path to the GFF3 annotation file.
    :type annotation: str
    :param region: Genomic region to visualize.
    :type region: str
    :param bed_path: Path to the BED file with sequence variants.
    :type bed_path: str
    :param bigwig_file: Path to the BigWig file containing read coverage data.
    :type bigwig_file: str
    :param bigwig_consensus_file: Path to the consensus BigWig file.
    :type bigwig_consensus_file: str
    :param output_dir: Directory where the plot and ini file will be saved.
    :type output_dir: str
    :param sample_name: Sample identifier.
    :type sample_name: str
    :param user_lab: User lab identifier.
    :type user_lab: str
    :return: File path of the generated coverage plot.
    :rtype: str
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    tracks_ini = os.path.join(output_dir, f"{user_lab}_{sample_name}_tracks.ini")
    plot_output = os.path.join(output_dir, f"{user_lab}_{sample_name}_coverage_plot.png")

    # Create tracks.ini configuration with three sections:
    # Reference, Annotation, and BigWig
    tracks=f"""
        [annotation]
        file = {annotation}
        title = {reference_genome.split("/")[-1].split(".")[0]} genes
        prefered_name = gene_name
        color = green
        style = UCSC
        height = 2
        file_type = gtf
        merge_transcripts = true

        [bed]
        file = {bed_path}
        title = Seq. variants
        color = purple
        height = 3
        file_type = bed

        [bigwig]
        file = {bigwig_consensus_file}
        title = Seq. coverage
        color = grey
        height = 1
        file_type = bigwig
        """
    # Only write the read coverage bigwig section if reads were uploaded
    if bigwig_file!=None:
            tracks+=f"""
            
        [bigwig]
        file = {bigwig_file}
        title = Read Coverage
        color = blue
        height = 3
        file_type = bigwig
        """
    with open(tracks_ini, "w") as ini:
        ini.write(tracks)

    # Run pyGenomeTracks to generate the plot for the given region
    command = [
        "pyGenomeTracks",
        "--tracks", tracks_ini,
        "--region", region,
        "--outFileName", plot_output,
        "--dpi", str(300)
    ]
    
    try:
        subprocess.run(command, check=True)
        return plot_output  # Return the path to the generated image
    except subprocess.CalledProcessError as e:
        print(f"Error generating plot: {e}")
        return None


def create_element(name):
    """
    Creates an XML element for use in DOCX formatting.

    :param name: The name of the XML element.
    :type name: str
    :return: The created XML element.
    :rtype: OxmlElement
    """
    return OxmlElement(name)

def create_attribute(element, name, value):
    """
    Adds an attribute to an XML element.

    :param element: The XML element to modify.
    :type element: OxmlElement
    :param name: The attribute name.
    :type name: str
    :param value: The attribute value.
    :type value: str
    """
    element.set(ns.qn(name), value)

def add_page_number(run):
    """
    Inserts a page number field in a DOCX document.

    :param run: A DOCX run object where the page number field is inserted.
    :type run: docx.text.run.Run
    """
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def generate_docx_report(report_data, base_dir, role, user_lab, distribution):
    """
    Generates a DOCX report summarizing genomic analysis results for a given distribution.

    The report includes sample statistics, visualizations, and formatted text.

    :param report_data: Processed data containing genomic viral analysis results.
    :type report_data: dict
    :param base_dir: Base directory containing report-related files.
    :type base_dir: str
    :param role: User role determining report formatting.
    :type role: str
    :param user_lab: User lab identifier.
    :type user_lab: str
    :param distribution: The distribution to .
    :type distribution: str
    :return: the generated DOCX report.
    :rtype: docx
    """
    report_data = process_all_reports(base_dir)
    distribution=str(base_dir.split("/")[1])
    samples_file = f"{base_dir}/samples.txt"
    sample_reference_map = {}
    with open(samples_file, "r") as f:
        for line in f:
            sample_id, reference = line.strip().split()
            sample_reference_map[sample_id] = reference

    from docx.oxml import OxmlElement, ns

    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        #This is only needed if you're using the builtin style above
    def get_or_create_hyperlink_style(d):
        """If this document had no hyperlinks so far, the builtin
        Hyperlink style will likely be missing and we need to add it.
        There's no predefined value, different Word versions
        define it differently.
        This version is how Word 2019 defines it in the
        default theme, excluding a theme reference.
        """
        if "Hyperlink" not in d.styles:
            if "Default Character Font" not in d.styles:
                ds = d.styles.add_style("Default Character Font",
                                        docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                        True)
                ds.element.set(docx.oxml.shared.qn('w:default'), "1")
                ds.priority = 1
                ds.hidden = True
                ds.unhide_when_used = True
                del ds
            hs = d.styles.add_style("Hyperlink",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            hs.base_style = d.styles["Default Character Font"]
            hs.unhide_when_used = True
            hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
            hs.font.underline = True
            del hs

        return "Hyperlink"

    def add_hyperlink(paragraph, text, url):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a new run object (a wrapper over a 'w:r' element)
        new_run = docx.text.run.Run(
            docx.oxml.shared.OxmlElement('w:r'), paragraph)
        new_run.text = text
        

        # Set the run's style to the builtin hyperlink style, defining it if necessary
        new_run.style = get_or_create_hyperlink_style(part.document)
        new_run.font.size = Pt(9)
        # Alternatively, set the run's formatting explicitly
        # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
        # new_run.font.underline = True

        # Join all the xml elements together
        hyperlink.append(new_run._element)
        paragraph._p.append(hyperlink)
        return hyperlink

    # Load the existing template
    doc = Document("project/static/templateV4.docx")
    table_count=1
    figure_count=1

    # Create or update Heading 1 style
    styles = doc.styles
    try:
        heading1 = styles['Heading 1']
    except KeyError:
        heading1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    heading1.font.name = 'Arial'
    heading1.font.size = Pt(14)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)  # Dark blue

    # Create or update Heading 2 style
    try:
        heading2 = styles['Heading 2']
    except KeyError:
        heading2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    heading2.font.name = 'Arial'
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(0x33, 0x66, 0x99)  # Medium blue

    # Create or update Heading 3 style
    try:
        heading3 = styles['Heading 3']
    except KeyError:
        heading3 = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    heading3.font.name = 'Arial'
    heading3.font.size = Pt(9)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(0x66, 0x99, 0xCC)  # Lighter blue

    # Access the header
    header = doc.sections[0].header

    # Retrieve the first table in the header (assumes your template already has one)
    table = header.tables[0]


    # --- Populate the content cells ---
    # Row 0 (content row)
    cell = table.cell(0, 2).paragraphs[0].add_run("WHO RSV Sequencing EQA")
    cell = table.cell(0, 4).paragraphs[0].add_run(f"Laboratory : {user_lab}")

    # Row 2 (content row)
    # Create a paragraph in the cell and set the run to bold for the distribution text.
    dist_cell = table.cell(2, 2)
    dist_paragraph = dist_cell.paragraphs[0]
    dist_paragraph.clear()  # Clear previous content if any
    run_label = dist_paragraph.add_run("Distribution : ")
    run_value = dist_paragraph.add_run(f"{distribution}")
    run_value.bold = True


    # Row 4 (content row)
    current_date = datetime.datetime.now().strftime("%d-%b-%Y")
    cell = table.cell(4, 2).paragraphs[0].add_run(f"Dispatch Date : 17-Jun-2024")#{current_date}

    # --- Set column widths ---
    table.autofit = False
    table.columns[0].width = Cm(4.71)  # Logo (merged) column
    table.columns[1].width = Cm(0.18)    # Blank separator
    table.columns[2].width = Cm(10.18)   # First content column
    table.columns[3].width = Cm(0.18)    # Blank separator
    table.columns[4].width = Cm(4.71)    # Second content column

    # --- Set row heights ---
    for row_idx in [0, 2, 4]:
        row = table.rows[row_idx]
        row.height = Cm(0.71)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for row_idx in [1, 3]:
        row = table.rows[row_idx]
        
        
        row.height = Cm(0.11)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # --- Optionally, adjust cell alignment and apply custom borders ---
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    def subtype_assignment(user_subtype, intended_subtype):
        if user_subtype=="original":
            return intended_subtype
        elif user_subtype=="alternative":
            return "RSV-B" if intended_subtype=="RSV-A" else "RSV-A"
        return user_subtype  

    sample_html_reports = {}

    for lab, samples in report_data.items():
        for sample_name, metrics in samples.items():
            if sample_name not in sample_html_reports:
                sample_html_reports[sample_name] = {}
            submission = (
                Submission.query
                .join(Organization)
                .filter(
                    Distribution.name == distribution,
                    Submission.sample == sample_name,
                    Organization.name == lab
                )
                .first()
            )
            if submission:
                seq_type = submission.sequencing_type
            else:
                seq_type = "N/A"
            sample_html_reports[sample_name][lab] = metrics
            sample_html_reports[sample_name][lab]["sequencing_platform"]=seq_type

    sample_html_reports = dict(sorted(sample_html_reports.items()))
    print(sample_html_reports)

    def compute_evaluation_data(sample_html_reports, sample_reference_map, user_lab):
        """Computes evaluation data from sample reports."""
        
        evaluation_data = {"RSV_Subtyping": [], "Clade":[],"Legacy_clade":[], "Genome Coverage (%)": [],"Ns in Sequence (%)": [], "Similarity (%)": [],"Read Coverage (mean)": []}
        for sample_name, lab_data in sample_html_reports.items():
            lab_count = 0
            lab_pass = 0
            reference_metrics = lab_data.get("9999", {})  # Extract reference metrics
            user_metrics = lab_data.get(user_lab, {})  # Extract user metrics
            if user_metrics=={}:
                continue
            #subtype
            intended_subtype = "RSV-B" if sample_reference_map[sample_name]=="EPI_ISL_1653999" else "RSV-A"
            print(sample_name)
            print("intended_subtype ", intended_subtype)  
            reference_subtype = subtype_assignment(reference_metrics['subtype'], intended_subtype)
            print("reference_metrics['subtype'] ", reference_metrics['subtype'])
            print("reference_subtype ", reference_subtype)
            user_subtype = subtype_assignment(user_metrics['subtype'], intended_subtype)
            print("user_metrics['subtype'] ", user_metrics['subtype'])
            print("user_subtype ", user_subtype)
            for lab, metrics in lab_data.items():
                if lab==user_lab or lab == "9999":
                    continue
                elif metrics["coverage"]=="N/A":
                    lab_count+=1
                else:
                    lab_count+=1
                    if metrics["subtype"]=="original":
                        lab_pass+=1
            subtyping=[sample_name, user_subtype, intended_subtype, reference_subtype, "Met" if user_subtype == intended_subtype else "Not met", f"{lab_pass}/{lab_count} ({(lab_pass * 100) // lab_count}%)"]
            evaluation_data["RSV_Subtyping"].append(subtyping)    
            
            #clade
            lab_count = 0
            lab_pass = 0
            intended_clade = reference_metrics["clade"]
            reference_clade = intended_clade
            user_clade = user_metrics["clade"]
            for lab, metrics in lab_data.items():
                if lab==user_lab or lab == "9999":
                    continue
                elif metrics["coverage"]=="N/A":
                    lab_count+=1
                else:
                    lab_count+=1
                    if metrics["clade"]==intended_clade:
                        lab_pass+=1
            clade=[sample_name, user_clade, intended_clade, reference_clade, "Met" if user_clade == intended_clade else "Not met", f"{lab_pass}/{lab_count} ({(lab_pass * 100) // lab_count}%)"]
            evaluation_data["Clade"].append(clade)  

            #g_clade
            lab_count = 0
            lab_pass = 0
            intended_clade = reference_metrics["G_clade"]
            reference_clade = intended_clade
            user_clade = user_metrics["G_clade"]
            for lab, metrics in lab_data.items():
                if lab==user_lab or lab == "9999":
                    continue
                else:
                    lab_count+=1
                    if metrics["G_clade"]==intended_clade:
                        lab_pass+=1
            g_clade=[sample_name, user_clade, intended_clade, reference_clade, "Met" if user_clade == intended_clade else "Not met", f"{lab_pass}/{lab_count} ({(lab_pass * 100) // lab_count}%)"]
            evaluation_data["Legacy_clade"].append(g_clade)

            def format_IQR_string(aggregated_metrics):
                if not aggregated_metrics:  # Handle empty list case
                    return "N/A"

                mean_value = round(np.mean(aggregated_metrics))  # Compute mean and round
                q1, q3 = np.percentile(aggregated_metrics, [25, 75])  # Compute Q1 and Q3
                q1, q3 = round(q1), round(q3)  # Round percentiles
                
                return f"{mean_value} ({q1}-{q3})"

            #genome_coverage (these four loops could be better refactored into one)
            lab_count = 0
            lab_pass = 0
            reference_coverage = reference_metrics["coverage"]*100
            user_coverage = user_metrics["coverage"]*100 
            aggregated_metrics = []
            for lab, metrics in lab_data.items():
                if lab==user_lab or lab == "9999":
                    continue
                elif metrics["coverage"]=="N/A":
                    lab_count+=1
                else:
                    print(metrics["coverage"])
                    lab_count+=1
                    aggregated_metrics.append(metrics["coverage"]*100)
                    if metrics["coverage"]>0.90:
                        lab_pass+=1
            coverage=[sample_name, round(user_coverage,1) if user_metrics["coverage"]!="N/A" else "N/A", "90% or higher", round(reference_coverage,1), "Met" if user_metrics["coverage"]!="N/A" and user_coverage >90 else "Not met", format_IQR_string(aggregated_metrics), f"{lab_pass}/{lab_count} ({(lab_pass * 100) // lab_count}%)"]
            evaluation_data["Genome Coverage (%)"].append(coverage)

            #Ns
            lab_count = 0
            lab_pass = 0
            reference_n = reference_metrics["Ns"]
            user_n = user_metrics["Ns"]
            aggregated_metrics = []
            for lab, metrics in lab_data.items():
                if lab==user_lab or lab == "9999":
                    continue
                elif metrics["coverage"]=="N/A":
                    lab_count+=1
                else:
                    lab_count+=1
                    aggregated_metrics.append(metrics["Ns"])
                    if metrics["Ns"]<=2:
                        lab_pass+=1
            ns=[sample_name, round(user_n,1) if user_n!="N/A" else "N/A", "2% or lower", round(reference_n,1), "Met" if user_n!="N/A" and user_n <=2 else "Not met", format_IQR_string(aggregated_metrics), f"{lab_pass}/{lab_count} ({(lab_pass * 100) // lab_count}%)"]
            evaluation_data["Ns in Sequence (%)"].append(ns)

            #similarity
            lab_count = 0
            lab_pass = 0
            reference_similarity = reference_metrics["similarity"]
            user_similarity = user_metrics["similarity"]
            aggregated_metrics = []
            for lab, metrics in lab_data.items():
                if lab==user_lab or lab == "9999":
                    continue
                elif metrics["coverage"]=="N/A":
                    lab_count+=1
                else:
                    lab_count+=1
                    aggregated_metrics.append(metrics["similarity"])
                    if metrics["similarity"]>95:
                        lab_pass+=1
            similarity=[sample_name, round(user_similarity,1) if user_similarity!="N/A" else "N/A", "95% or higher", round(reference_similarity,1), "Met" if user_similarity!="N/A" and user_similarity >95 else "Not met", format_IQR_string(aggregated_metrics), f"{lab_pass}/{lab_count} ({(lab_pass * 100) // lab_count}%)"]
            evaluation_data["Similarity (%)"].append(similarity)

            #Mean coverage depth
            lab_count = 0
            lab_pass = 0
            reference_coverage = reference_metrics["Mean coverage depth"]
            user_coverage = round(user_metrics["Mean coverage depth"],0) if user_metrics["Mean coverage depth"]!="N/A" else "N/A"
            aggregated_metrics = []
            for lab, metrics in lab_data.items():
                if lab==user_lab or lab == "9999":
                    continue
                elif metrics["Mean coverage depth"]=="N/A":
                    lab_count+=1
                else:
                    lab_count+=1
                    aggregated_metrics.append(metrics["Mean coverage depth"])
                    if metrics["Mean coverage depth"]>50:
                        lab_pass+=1
            coverage=[sample_name, user_coverage, "50 or higher", round(reference_coverage,0), "Met" if user_coverage!="N/A" and user_coverage >50 else "Not met", format_IQR_string(aggregated_metrics), f"{lab_pass}/{lab_count} ({(lab_pass * 100) // lab_count}%)"]
            evaluation_data["Read Coverage (mean)"].append(coverage)
        return evaluation_data


    def add_evaluation_tables(doc, evaluation_data, run_id="WR024", role="user"):
        """Adds RSV Evaluation and Sequencing Quality tables to the DOCX report with transposed format."""

        def apply_color(run, text):
            """Apply color coding for pass/fail."""
            run.text = text
            if text.lower() == "Met":
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif text.lower() == "Not met":
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red

        def apply_background_color(cell, color_rgb):
            """Apply background color to a table cell."""
            # Create the XML shading element and apply it to the cell
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_rgb)))

        def apply_font_size(cell, size_pt):
            """Apply a smaller font size to a table cell."""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size_pt)

        
        if role=="user":
            # Table Summary: RSV Subtyping and Clade Assignment
            table1 = doc.add_table(rows=1, cols=7)  # 7 Columns: Indicator, Specimen ID, Your result, Intended, Reference, Score, Participants
            table1.style = "Table Grid"
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Table 1 Header
            hdr_cells = table1.rows[0].cells
            headers = ["Indicator", "Specimen ID", "Your result", "Intended Result", "Reference Lab result", "Your score", "Participant with intended results"]
        else:
            # Table Summary: RSV Subtyping and Clade Assignment
            table1 = doc.add_table(rows=1, cols=5)
            table1.style = "Table Grid"
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Table 1 Header
            hdr_cells = table1.rows[0].cells
            headers = ["Indicator", "Specimen ID", "Intended Result", "Reference Lab result", "Participant with intended results"]
        for i, text in enumerate(headers):
            hdr_cells[i].text = text
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)  # Slightly smaller font size

        # Apply background color to the header row (light grey)
        for cell in hdr_cells:
            apply_background_color(cell, "D3D3D3")  # Light grey background

        # Populate Table 1 with merged Indicator cells and color rows differently
        indicator_colors = {
            "RSV_Subtyping": "FFDDC1",  # Light pink
            "Clade": "D1E8E2",  # Light teal
            "Legacy_clade": "E2D1E8",  # Light purple
        }

        for indicator in ["RSV_Subtyping", "Clade", "Legacy_clade"]:
            rows = evaluation_data[indicator]
            first_cell = None
            row_color = indicator_colors.get(indicator, "FFFFFF")  # Default to white if not found

            for i, row in enumerate(rows):
                row_cells = table1.add_row().cells
                if i == 0:
                    first_cell = row_cells[0]  # Store first cell for merging
                    if indicator=="RSV_Subtyping":
                        first_cell.text = "RSV Subtyping"
                    elif indicator=="Clade":
                        first_cell.text = "Lineage"
                    elif indicator=="Legacy_clade":
                        first_cell.text = "Legacy lineage"
                else:
                    row_cells[0]._tc.merge(first_cell._tc)  # Merge vertically

                row_cells[1].text = str(row[0])
                if role=="user":  
                    row_cells[2].text = row[1] if row[1] else ""  
                    row_cells[3].text = row[2] if row[2] else ""  
                    row_cells[4].text = row[3] if row[3] else ""  
                    apply_color(row_cells[5].paragraphs[0].add_run(), row[4])  
                    row_cells[6].text = row[5]
                else:
                    row_cells[2].text = row[1] if row[1] else ""    
                    row_cells[3].text = row[3] if row[3] else ""   
                    row_cells[4].text = row[5]

                # Apply background color and smaller font size to each row under the current indicator
                for cell in row_cells:
                    apply_background_color(cell, row_color)  # Apply unique color per indicator
                    apply_font_size(cell, 9)  # Set font size to 6 pt for the row
        doc.add_paragraph()
        doc.add_paragraph()
            # Create a table with one row and one cell
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Apply a built‑in style that has a border (e.g., 'Table Grid')
        table.style = "Table Grid"

        # Remove extra cell margins if desired (requires modifying underlying XML)
        cell = table.rows[0].cells[0]
        cell.paragraphs[0].add_run("Purpose of the EQA:\n").bold=True
        cell.paragraphs[0].add_run(
            "   •   Assess the accuracy of RSV sequencing:\n"
            "           ◦ RSV subtype and lineage by comparing sequence to the reference GISAID sequences using ")
        add_hyperlink(cell.paragraphs[0],"Nextclade","https://docs.nextstrain.org/projects/nextclade/en/stable/user/nextclade-cli/index.html")
        cell.paragraphs[0].add_run(
            ".\n"
            "   •   Measure the quality of viral sequencing.\n"
            "           ◦ Genome Coverage, Ns in sequence and Sequence Similarity were assessed following the procedure above.\n"
            "           ◦ Read Coverage at GISAID reference sequence's genomic positions was determined using ")
        add_hyperlink(cell.paragraphs[0],"QualiMap2","http://qualimap.conesalab.org/")
        cell.paragraphs[0].add_run(
            ".\n\n"
            "Appendix 1 provides a summary of the procedures for specimen preparation, data submission, "
            "and analysis, along with details on result validation, quality metrics, and laboratory compliance.\n\n"
            "Specimens for this EQA were distributed by UK NEQAS Microbiology as part of the WHO Molecular Detection of RSV "
            "Distribution 5791. Specimens with detectable virus are either sequenced in-house or sent to a reference laboratory "
            "following routine procedures.\n\n"
            "As part of the sequencing result submission, participants complete a survey on sequencing technology. "
            "FASTA, FASTQ and/or BAM files are evaluated for sequencing quality metrics, including read coverage and accuracy "
            "based on the comparison to GISAID reference sequences EPI_ISL_412866 (RSV A) or EPI_ISL_1653999 (RSV B) (see Appendix 1 for definitions).\n"
            "Each participant receives a report outlining the comparison to:\n"
            "     •   One of the  reference sequence in GISAID (EPI_ISL_412866 (RSV A) or EPI_ISL_1653999 (RSV B)).\n" 
            "     •   A reference lab chosen by UK NEQAS which sequenced the distributed samples.\n"
            "     •   The participants' results are also compared to the mean of the aggregated results submitted by other participants.\n"
            "Additionally, the comparison to GISAID reference sequences enables lineage assignment of the submitted sequence data "
            "using ")
        add_hyperlink(cell.paragraphs[0],"Nextclade","https://clades.nextstrain.org/dataset")
        cell.paragraphs[0].add_run(".\n\nThe Figures and Tables on page 3 and beyond display the results for various Quality Metrics. "
            "This data is provided for your reference only and is not included in your quality assessment exercise. "
            "The most commonly used methods, along with the specific method(s) used in your laboratory, marked with an arrow for easy identification are displayed. \n\n"
        )

        # Optionally, modify the font of the text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9)

        doc.add_page_break()

        # Table 2: Sequencing Quality
        doc.add_heading("Table 1: Sequencing Quality", level=2)

        if role=="user":
            table2 = doc.add_table(rows=1, cols=8)  # 8 Columns: Indicator, Specimen ID, Your result, Recommended, Reference, Score, Mean (IQR), Participants meeting threshold
            table2.style = "Table Grid"
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Create the first header row
            table2_header = table2.rows[0].cells
            headers = ["Indicator", "Specimen ID", "Your result", "Recommended Value*", "Reference Lab result", "Your score", "Participant summary"]
        else:
            table2 = doc.add_table(rows=1, cols=6) 
            table2.style = "Table Grid"
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Create the first header row
            table2_header = table2.rows[0].cells
            headers = ["Indicator", "Specimen ID", "Recommended Value*", "Reference Lab result", "Participant summary"]
        for i, text in enumerate(headers):
            table2_header[i].text = text
            table2_header[i].paragraphs[0].runs[0].bold = True
            table2_header[i].paragraphs[0].runs[0].font.size = Pt(9)  # Slightly smaller font size
        if role=="user": 
            table2_header[6].merge(table2_header[7])
        else:
            table2_header[4].merge(table2_header[5])

        # Apply background color to the header row (light grey)
        for cell in table2_header:
            apply_background_color(cell, "D3D3D3")  # Light grey background

        # Add the second empty row
        empty_row = table2.add_row().cells

        # Merge all cells except the one under "Participant Summary" (column 6)
        if role=="user":
            indexCol=6
        else:
            indexCol=4
        for i in range(indexCol):  # Merge the first 6 columns
            empty_row[i].merge(table2_header[i])

        
        # Split the "Participant Summary" cell into two: "Mean (IQR)" and "Participants Meeting Threshold"
        empty_row[indexCol].text = "Mean (IQR)"
        empty_row[indexCol+1].text = "Participants meeting threshold"

        # Optional: Bold the new header cells and make their font smaller
        empty_row[indexCol].paragraphs[0].runs[0].font.size = Pt(9)  # Smaller font size
        empty_row[indexCol+1].paragraphs[0].runs[0].font.size = Pt(9)  # Smaller font size

        # Apply background color to the empty row (light blue)
        for cell in empty_row:
            apply_background_color(cell, "ADD8E6")  # Light blue background

        # Populate Table 2 with merged Indicator cells and color rows differently
        indicator_colors_table2 = {
            "Genome Coverage (%)": "FFDDC1",  # Light pink
            "Ns in Sequence (%)": "D1E8E2",  # Light teal
            "Similarity (%)": "E2D1E8",  # Light purple
            "Read Coverage (mean)": "FFFACD",  # Light yellow
        }

        for indicator in ["Genome Coverage (%)", "Ns in Sequence (%)", "Similarity (%)", "Read Coverage (mean)"]:
            rows = evaluation_data[indicator]
            first_cell = None
            row_color = indicator_colors_table2.get(indicator, "FFFFFF")  # Default to white if not found

            for i, row in enumerate(rows):
                row_cells = table2.add_row().cells
                if i == 0:
                    first_cell = row_cells[0]  # Store first cell for merging
                    first_cell.text = indicator
                else:
                    row_cells[0]._tc.merge(first_cell._tc)  # Merge vertically
    
                row_cells[1].text = str(row[0])  
                if role=="user":
                    row_cells[2].text = str(row[1])  
                    row_cells[3].text = row[2]  
                    row_cells[4].text = str(row[3])  
                    apply_color(row_cells[5].paragraphs[0].add_run(), row[4])  
                    row_cells[6].text = row[5]  
                    row_cells[7].text = row[6]  
                else: 
                    row_cells[2].text = row[2]  
                    row_cells[3].text = str(row[3])   
                    row_cells[4].text = row[5]  
                    row_cells[5].text = row[6]

                # Apply background color and smaller font size to each row under the current indicator
                for cell in row_cells:
                    apply_background_color(cell, row_color)  # Apply unique color per indicator
                    apply_font_size(cell, 9)  # Set font size to 6 pt for the row
        
        # **Adding the Sequencing Quality Notes with Smaller Font**
        para = doc.add_paragraph("\n* Recommended Value")

        bullet_points = [
            "Obtained sufficient genome coverage (90% or higher).",
            "Maintained Ns in Sequence within acceptable limits (2% or lower).",
            "Obtained sufficient Similarity (95% or higher).",
            "Obtained sufficient Read Coverage (Mean depth of 50 or higher)."
        ]
        for point in bullet_points:
            para = doc.add_paragraph(f"{point}", style="List Bullet")
            para.runs[0].font.size = Pt(8)  # Set font size to 6 pt
        
        doc.add_page_break()


        return doc


    evaluation_data = compute_evaluation_data(sample_html_reports, sample_reference_map, user_lab)
    doc = add_evaluation_tables(doc, evaluation_data, user_lab, role)

    # If the user is not a superuser, filter and aggregate data
    others_label=""
    agg_users={}
    sample_html_reports_original=copy.deepcopy(sample_html_reports)
    if not role == "superuser":
        processed_reports = {}
        for sample_name, labs_data in sample_html_reports.items():
            if user_lab not in labs_data:
                continue  # Skip samples the user has no access to

            user_metrics = labs_data[user_lab]
            reference_metrics = labs_data.get("9999", {})  # Extract reference metrics
            subtype_counts = {}
            clade_counts = {}
            g_clade_counts = {}

            aggregated_metrics = {
                "coverage": 0,
                "Ns": 0,
                "similarity": 0,
                "Mean coverage depth": 0,
                "lab_count": 0,
                "clade": "",
                "G_clade": "",
                "subtype":""
            }

            # Aggregate data from other labs
            for lab, metrics in labs_data.items():
                if lab != user_lab and lab != "9999" and metrics["coverage"]!="N/A":  # Exclude user lab and reference lab
                    aggregated_metrics["coverage"] += metrics["coverage"]
                    aggregated_metrics["Ns"] += metrics["Ns"]
                    aggregated_metrics["similarity"] += metrics["similarity"]
                    if metrics["Mean coverage depth"]!="N/A":
                        aggregated_metrics["Mean coverage depth"] += metrics["Mean coverage depth"]
                    aggregated_metrics["lab_count"] += 1

                    # Count clade occurrences
                    clade = metrics.get("clade", "")
                    subtype = metrics.get("subtype", "")
                    g_clade = metrics.get("G_clade", "")
                    if subtype:
                        subtype_counts[subtype] = subtype_counts.get(subtype, 0) + 1
                    if clade:
                        clade_counts[clade] = clade_counts.get(clade, 0) + 1
                    if g_clade:
                        g_clade_counts[g_clade] = g_clade_counts.get(g_clade, 0) + 1

            # Calculate averages
            if aggregated_metrics["lab_count"] > 0:
                aggregated_metrics["coverage"] = round(
                    (aggregated_metrics["coverage"] / aggregated_metrics["lab_count"]), 2
                )
                aggregated_metrics["Ns"] = round(
                    aggregated_metrics["Ns"] / aggregated_metrics["lab_count"], 2
                )
                aggregated_metrics["similarity"] = round(
                    aggregated_metrics["similarity"] / aggregated_metrics["lab_count"], 2
                )
                aggregated_metrics["Mean coverage depth"] = round(
                    aggregated_metrics["Mean coverage depth"] / aggregated_metrics["lab_count"], 2
                )

            # Determine the most frequent clades
            if subtype_counts:
                aggregated_metrics["subtype"] = max(subtype_counts, key=subtype_counts.get)
            if clade_counts:
                aggregated_metrics["clade"] = max(clade_counts, key=clade_counts.get)
            if g_clade_counts:
                aggregated_metrics["G_clade"] = max(g_clade_counts, key=g_clade_counts.get)

            # Prepare processed data for the user
            others_label="Others ("+str(aggregated_metrics["lab_count"])+")"
            processed_reports[sample_name] = {
                user_lab: user_metrics,
                others_label: aggregated_metrics,
                "Reference": reference_metrics
            }
            agg_users[sample_name]=aggregated_metrics

        sample_html_reports = processed_reports


    sample_html_reports = dict(sorted(sample_html_reports.items()))
    
    # Handle lab anonymization based on the user's role
    if role == 'superuser':
        # Iterate through each sample's data and anonymize lab names
        for sample, data in sample_html_reports.items():
            count = 1
            anonymized_data = {}
            for i, (lab, metrics) in enumerate(data.items()):
                '''if lab == '9999' or lab == 'Reference':
                    anonymized_data['9999'] = metrics  # Keep the reference lab as '9999'
                else:
                    anonymized_data[f'{count:01d}'] = metrics  # Anonymize labs as single-digit numbers
                    count += 1'''
                anonymized_data[lab] = metrics
            # Replace the sample data with anonymized data
            sample_html_reports[sample] = anonymized_data

    # Sort sample_html_reports such that user_lab comes first, reference second
    for sample, data in sample_html_reports.items():
        sorted_data = {
            key: data[key]
            for key in sorted(
            data,
            key=lambda k: (
                0 if k == user_lab else (1 if (k.lower() == "reference" or k=="9999") else 2),
                k
            )
        )
        }
        sample_html_reports[sample] = sorted_data

    # Sort the sample_html_reports dictionary by sample name
    sample_html_reports = dict(sorted(sample_html_reports.items()))


    # Add sample plots and tables to the DOCX file
    for sample, data in sample_html_reports.items():
        #plot_path, read_plot_path = generate_two_plots(sample, data, role)  # Generate and get the plot path
        platform_plot_path, read_platform_plot_path, user_platform = generate_aggregated_plot_by_platform(sample, sample_html_reports_original[sample], user_lab, role)
        intended_subtype = "RSV-B" if sample_reference_map[sample]=="EPI_ISL_1653999" else "RSV-A"
        doc.add_heading(f'Sample {sample}', level=1)
        doc.add_heading(f'Lineage assignment', level=2)

        # Extract Clade and G_clade assignments for the 'reference'
        reference_clade = None
        reference_g_clade = None

        # Iterate through data to find the reference clade and g_clade
        for lab, metrics in data.items():
            if lab == '9999':  # Reference participant
                reference_clade = metrics['clade']
                reference_g_clade = metrics['G_clade']
                break
            elif lab== 'Reference':
                reference_subtype = subtype_assignment(metrics['subtype'], intended_subtype)
                reference_clade = metrics['clade']
                reference_g_clade = metrics['G_clade']
                break

        # Add an introductory paragraph about clade assignments
        if role=='superuser':
            #doc.add_paragraph("All participants matched RSV subtype assignments with reference lab.", style='List Bullet')
            mistakes_text = "All participants matched lineage assignments with reference lab."  # Default
            for lab, metrics in data.items():
                if lab != '9999' and lab!='Reference':  # Exclude the reference itself
                    if metrics['clade'] != reference_clade or metrics['G_clade'] != reference_g_clade:
                        mistakes_text = "Some participants' clade assignments differed from the reference lab."
                        mistakes_found = True
                        break
            doc.add_paragraph(mistakes_text, style='List Bullet')
        else:
            mistakes_text = "Your lab's lineage assignment matches the reference lab's."
            if data[user_lab]['clade'] != reference_clade:
                mistakes_text = "Your lab's lineage assignment does not match the reference lab's."
            doc.add_paragraph(mistakes_text, style='List Bullet')
        
        # Caption for the Clade table
        clade_caption = doc.add_paragraph()
        clade_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        clade_runner = clade_caption.add_run(f"\nTable {str(table_count)}. Lineage assignments for sample {sample}, including RSV subtype.")
        clade_runner.bold = True
        clade_runner.italic = True
        table_count += 1  # Increment the table count

        # Add the Clade and G_clade table with RSV subtype
        clade_table = doc.add_table(rows=1, cols=4)  # Creating a table with 4 columns (Participant, Subtype, Clade, Legacy clade)
        clade_table.style = "Table Grid"

        # Adding the header row for the table
        clade_hdr_cells = clade_table.rows[0].cells
        clade_hdr_cells[0].text = 'Participant'
        clade_hdr_cells[0].paragraphs[0].runs[0].bold = True
        clade_hdr_cells[1].text = 'RSV Subtype'
        clade_hdr_cells[1].paragraphs[0].runs[0].bold = True
        clade_hdr_cells[2].text = 'Lineage'
        clade_hdr_cells[2].paragraphs[0].runs[0].bold = True
        clade_hdr_cells[3].text = 'Legacy lineage'
        clade_hdr_cells[3].paragraphs[0].runs[0].bold = True

        # Adding data rows for each participant and highlighting mismatches
        labCount = 0
        for lab, metrics in data.items():
            labName=lab if lab!='Reference' else 'Reference Lab (NEQAS)'

            # Add a new row to the table
            row_cells = clade_table.add_row().cells

            # Participant column
            row_cells[0].text = labName
            row_cells[0].paragraphs[0].runs[0].italic = True  # Italicize the first column (Participant)

            # RSV Subtype column
            subtype_cell = row_cells[1]
            subtype_cell.text = subtype_assignment(metrics['subtype'],intended_subtype)
            # Highlight mismatched Clade values
            #if metrics['subtype'] == "alternative":
            #    subtype_cell.paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

            # Clade column
            clade_cell = row_cells[2]
            clade_cell.text = metrics['clade']
            # Highlight mismatched Clade values
            #if metrics['clade'] != reference_clade:
            #    clade_cell.paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

            # Legacy clade column
            g_clade_cell = row_cells[3]
            g_clade_cell.text = metrics['G_clade']
            # Highlight mismatched Legacy clade values
            #if metrics['G_clade'] != reference_g_clade:
            #    g_clade_cell.paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
            if labName==user_lab:
                for i in range(4):
                    shading_elm = parse_xml(r'<w:shd {} w:fill="EFFA75"/>'.format(nsdecls('w')))
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)


        doc.add_heading(f'\nSequencing quality', level=2)

                # Initialize flags to check if any participants fail to meet thresholds
        failed_coverage = []
        failed_ns = []
        failed_similarity = []
        failed_coverage_depth = []

        # Initialize user-specific flags if the role is not a superuser
        user_failed_coverage = False
        user_failed_ns = False
        user_failed_similarity = False
        user_failed_coverage_depth = False

        # Loop through the data to check conditions and generate summary statements
        for lab, metrics in data.items():
            labName = lab

            # Check thresholds for all participants if the role is superuser
            if role == "superuser":
                if metrics['coverage'] == 'N/A' or metrics['coverage'] < 0.90:
                    failed_coverage.append(labName)
                if metrics['coverage'] == 'N/A' or metrics['Ns'] > 2:
                    failed_ns.append(labName)
                if metrics['coverage'] == 'N/A' or metrics['similarity'] < 95:
                    failed_similarity.append(labName)
                if metrics['Mean coverage depth'] == 'N/A' or float(metrics['Mean coverage depth']) < 50:
                    failed_coverage_depth.append(labName)

            # Check thresholds for user lab only if the role is not a superuser
            elif lab == user_lab:
                user_coverage='N/A' if metrics['coverage']=='N/A' else str(round(metrics['coverage'],2))
                user_ns='N/A' if metrics['Ns']=='N/A' else str(round(metrics['Ns'],2))
                user_similarity='N/A' if metrics['similarity']=='N/A' else str(round(metrics['similarity'],2))
                user_read_coverage='N/A' if metrics['Mean coverage depth']=='N/A' else str(round(metrics['Mean coverage depth'],2))
                if metrics['coverage'] == 'N/A' or metrics['coverage'] < 0.90:
                    user_failed_coverage = True
                if metrics['coverage'] == 'N/A' or metrics['Ns'] > 2:
                    user_failed_ns = True
                if metrics['coverage'] == 'N/A' or metrics['similarity'] < 95:
                    user_failed_similarity = True
                if metrics['Mean coverage depth'] == 'N/A' or float(metrics['Mean coverage depth']) < 50:
                    user_failed_coverage_depth = True

        # Generate feedback based on role
        if role == "superuser":
            # Add summary lines for superuser
            if failed_coverage:
                doc.add_paragraph(f"Participants {', '.join(failed_coverage)} failed to satisfy the threshold for genome coverage (90%).", style='List Bullet')
            else:
                doc.add_paragraph("All participants obtained sufficient genome coverage (90% or higher).", style='List Bullet')

            if failed_ns:
                doc.add_paragraph(f"Participants {', '.join(failed_ns)} failed to satisfy the threshold for Ns in Sequence (greater than 2%).", style='List Bullet')
            else:
                doc.add_paragraph("All participants maintained Ns in Sequence within acceptable limits (2% or lower).", style='List Bullet')

            if failed_similarity:
                doc.add_paragraph(f"Participants {', '.join(failed_similarity)} failed to satisfy the threshold for Similarity (95%).", style='List Bullet')
            else:
                doc.add_paragraph("All participants obtained sufficient Similarity (95% or higher).", style='List Bullet')

            if failed_coverage_depth:
                doc.add_paragraph(f"Participants {', '.join(failed_coverage_depth)} failed to satisfy the threshold for Read Coverage (Mean depth of 50).", style='List Bullet')
            else:
                doc.add_paragraph("All participants obtained sufficient Read Coverage (Mean depth of 50 or higher).", style='List Bullet')

        else:
            # Add summary lines for non-superuser
            if user_failed_coverage:
                doc.add_paragraph("Failed to satisfy the threshold for genome coverage (90%).", style='List Bullet')
            else:
                doc.add_paragraph("Obtained sufficient genome coverage (90% or higher).", style='List Bullet')

            if user_failed_ns:
                doc.add_paragraph("Failed to satisfy the threshold for Ns in Sequence (greater than 2%).", style='List Bullet')
            else:
                doc.add_paragraph("Obtained sufficient Ns in Sequence (2% or lower).", style='List Bullet')

            if user_failed_similarity:
                doc.add_paragraph("Failed to satisfy the threshold for Similarity (95%).", style='List Bullet')
            else:
                doc.add_paragraph("Obtained sufficient Similarity (95% or higher).", style='List Bullet')

            if user_failed_coverage_depth:
                doc.add_paragraph("Failed to satisfy the threshold for Read Coverage (Mean depth of 50).", style='List Bullet')
            else:
                doc.add_paragraph("Obtained sufficient Read Coverage (Mean depth of 50 or higher).", style='List Bullet')
        
        '''# Add the plot image to DOCX with a caption
        last_paragraph = doc.paragraphs[-1] 
        runner = last_paragraph.add_run(f"\n\n")
        doc.add_picture(plot_path, width=Inches(7.5))  # Adjust size as needed
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_paragraph()
        runner = para.add_run(f"Figure {str(figure_count)}. Quality metrics for sample {sample} for participants that submitted appropriate data files.\n\n")
        runner.bold = True
        runner.italic = True
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        figure_count+=1

        # Add the read coverage plot image to DOCX with a caption
        doc.add_picture(read_plot_path, width=Inches(7.5))  # Adjust size as needed
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_paragraph()
        runner = para.add_run(f"Figure {str(figure_count)}. Read coverages for sample {sample} for participants that submitted appropriate data files.\n\n\n")
        runner.bold = True
        runner.italic = True
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        figure_count+=1'''
        
        # Caption for the table
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runner = para.add_run(f"\nTable {str(table_count)}. Quality metrics data for sample {sample}")
        runner.bold = True
        runner.italic = True

        # Generate and insert the metrics table for the sample
        table = doc.add_table(rows=1, cols=5)  # Creating a table with 5 columns for metrics
        table.style = "Table Grid"

        # Adding the header row for the table
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Participant'
        hdr_cells[0].paragraphs[0].runs[0].bold = True
        hdr_cells[1].text = 'Genome Coverage (%)'
        hdr_cells[1].paragraphs[0].runs[0].bold = True
        hdr_cells[2].text = 'Ns in Sequence (%)'
        hdr_cells[2].paragraphs[0].runs[0].bold = True
        hdr_cells[3].text = 'Similarity (%)'
        hdr_cells[3].paragraphs[0].runs[0].bold = True
        hdr_cells[4].text = 'Read Coverage (Mean)'
        hdr_cells[4].paragraphs[0].runs[0].bold = True

        # Add data rows for each lab/sample entry
        labCount = 0
        for lab, metrics in data.items():
            #if lab == 'WR024':
            #    labName = 'reference'
            #else:
            #labCount += 1
            #labName = str(labCount)
            labName=lab if lab!='Reference' else 'Reference Lab (NEQAS)'

            # Add a new row to the table
            row_cells = table.add_row().cells

            # Fill in the cells for each column
            row_cells[0].text = labName  # First column with lab name (bold)
            row_cells[0].paragraphs[0].runs[0].italic = True  # Bold the first column

            # Coverage column (highlight if less than 90%)
            coverage_value = metrics['coverage'] * 100 if metrics['coverage']!="N/A" else "N/A"
            if metrics['coverage']!="N/A":
                row_cells[1].text = f"{coverage_value:.1f}"
                #if coverage_value < 90:
                #    row_cells[1].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW  # Highlight if less than 90%
                # Ns column (highlight if higher than 2%)
                ns_value = metrics['Ns']
                row_cells[2].text = f"{ns_value:.1f}"
                #if ns_value > 2:
                #    row_cells[2].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW  # Highlight if higher than 2%

                # Similarity column (highlight if less than 95%)
                similarity_value = metrics['similarity']
                row_cells[3].text = f"{similarity_value:.1f}"
                #if similarity_value < 95:
                #    row_cells[3].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW  # Highlight if less than 95%
            else:
                merged_cell = row_cells[1].merge(row_cells[2]).merge(row_cells[3])
                paragraph = merged_cell.paragraphs[0]
                run = paragraph.add_run("Nextclade error (low quality sequence)")
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Mean coverage depth column
            mean_coverage_depth_value = metrics['Mean coverage depth']
            if metrics['Mean coverage depth']!="N/A":
                row_cells[4].text = f"{mean_coverage_depth_value:.1f}"
            else:
                row_cells[4].text ="N/A"

            if labName==user_lab:
                for i in range(5):
                    shading_elm = parse_xml(r'<w:shd {} w:fill="EFFA75"/>'.format(nsdecls('w')))
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        # Increment figure and table count
        table_count += 1


        #genome tracks
        if role != "superuser" and os.path.isfile(os.path.join(f"data/{distribution}/{user_lab}/{sample}", f"{user_lab}_{sample}.bw")):
            doc.add_paragraph("\n\n")
            reference_genome = "project/static/genomes/EPI_ISL_412866/EPI_ISL_412866.fasta" if sample_reference_map[sample]=="EPI_ISL_412866" else "project/static/genomes/EPI_ISL_1653999/EPI_ISL_1653999.fasta"
            region="EPI_ISL_412866:1-15225" if sample_reference_map[sample]=="EPI_ISL_412866" else "EPI_ISL_1653999:1-15222"
            annotation = "project/static/genomes/EPI_ISL_412866/EPI_ISL_412866.gtf" if sample_reference_map[sample]=="EPI_ISL_412866" else "project/static/genomes/EPI_ISL_1653999/EPI_ISL_1653999.gtf"
            output_dir = "project/static/plots"
            bigwig_file_path = os.path.join(f"data/{distribution}/{user_lab}/{sample}", f"{user_lab}_{sample}.bw")
            if os.path.exists(bigwig_file_path):
                bigwig_copy=os.path.join(output_dir,f"{user_lab}_{sample}.bw")
                shutil.copy(bigwig_file_path,bigwig_copy)
            else:
                bigwig_copy=None
            bigwig_consensus_file_path = os.path.join(f"data/{distribution}/{user_lab}/{sample}", f"{user_lab}_{sample}_consensus.bw")
            bigwig_consensus_copy=os.path.join(output_dir,f"{user_lab}_{sample}_consensus.bw")

            shutil.copy(bigwig_consensus_file_path,bigwig_consensus_copy)
            bed_path = os.path.join(f"data/{distribution}/{user_lab}/{sample}", f"{user_lab}_{sample}_mutations.bed")
            bed_path_copy=os.path.join(output_dir,f"{user_lab}_{sample}_mutations.bed")
            shutil.copy(bed_path,bed_path_copy)
            output_dir = "project/static/plots"
            plot_path = create_pygenometracks_plot( reference_genome, annotation, region, bed_path_copy, bigwig_copy, bigwig_consensus_copy, output_dir, sample, user_lab)
            
            # Add the plot image to DOCX with a caption
            doc.add_picture(plot_path, width=Inches(6.5))  # Adjust size as needed
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runner = para.add_run(f"Figure {str(figure_count)}. Genomic visualisation of submitted sequence and reads.\n")
            runner.bold = True
            runner.italic = True
            figure_count+=1
        
        last_paragraph = doc.add_paragraph() 
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        last_paragraph.add_run().add_break(WD_BREAK.PAGE)
        doc.add_heading(f'Sequencing platforms\n', level=2)

        # Add the plot image to DOCX with a caption
        doc.add_picture(platform_plot_path, width=Inches(6.5))  # Adjust size as needed
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        last_paragraph = doc.paragraphs[-1]
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runner = para.add_run(f"Figure {str(figure_count)}. Quality metrics per sequencing platform.\n\n")
        runner.bold = True
        runner.italic = True
        figure_count+=1

        # Add the read coverage plot image to DOCX with a caption
        doc.add_picture(read_platform_plot_path, width=Inches(6.5))  # Adjust size as needed
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        last_paragraph = doc.paragraphs[-1]
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        runner = para.add_run(f"Figure {str(figure_count)}. Read coverage per sequencing platform.\n\n\n")
        runner.bold = True
        runner.italic = True
        figure_count+=1
        para.add_run().add_break(WD_BREAK.PAGE)

        def aggregate_sample_data(sample_data, user_platforms):
            """
            Aggregates sample metrics by sequencing platform.
            
            :param sample_data: Dictionary where each key is a lab and each value is a dict with metrics,
                                including 'coverage', 'similarity', 'Ns', 'Mean coverage depth',
                                and 'sequencing_platform'.
            :return: A dictionary keyed by sequencing platform with averaged metric values.
            """
            platform_data = {}
            for lab, metrics in sample_data.items():
                # Get sequencing platforms; default to "Unknown" if not provided.
                platforms = metrics.get("sequencing_platform", "Unknown").split(",")
                for platform in platforms:
                    platform = platform.strip()
                    if platform not in platform_data:
                        platform_data[platform] = {"coverage": [], "similarity": [], "Ns": [], "Mean coverage depth": []}
                    # Multiply coverage by 100 to get percentage, if available.
                    if metrics.get("coverage") != "N/A":
                        platform_data[platform]["coverage"].append(metrics["coverage"] * 100)
                    else:
                        platform_data[platform]["coverage"].append(0)
                    # Similarity, Ns and Mean coverage depth
                    if metrics.get("similarity") != "N/A":
                        platform_data[platform]["similarity"].append(metrics["similarity"])
                    else:
                        platform_data[platform]["similarity"].append(0)
                    if metrics.get("Ns") != "N/A":
                        platform_data[platform]["Ns"].append(metrics["Ns"])
                    else:
                        platform_data[platform]["Ns"].append(0)
                    if metrics.get("Mean coverage depth") != "N/A":
                        platform_data[platform]["Mean coverage depth"].append(metrics["Mean coverage depth"])
            
            # Compute averages for each platform
            aggregated = {}
            for platform, lists in platform_data.items():
                count = len(lists["coverage"])
                platformName=f"{platform} ({count})"
                aggregated[platformName] = {
                    "coverage": np.mean(lists["coverage"]) if lists["coverage"] else 0,
                    "similarity": np.mean(lists["similarity"]) if lists["similarity"] else 0,
                    "Ns": np.mean(lists["Ns"]) if lists["Ns"] else 0,
                    "Mean coverage depth": np.mean(lists["Mean coverage depth"]) if lists["Mean coverage depth"] else "N/A"
                }
            aggregatedFinal = {
                key: aggregated[key]
                for key in sorted(
                aggregated,
                key=lambda k: (
                    0 if any(user_p in k for user_p in user_platforms) else 1,  # User platforms first
                    k
                )
            )
            }
            return aggregatedFinal

        def create_platform_table(doc, aggregated_data, user_platforms):
            """
            Creates a DOCX table with aggregated quality metrics per sequencing platform.
            
            :param doc: The Document object to add the table to.
            :param aggregated_data: Dictionary with sequencing platform as key and averaged metrics as values.
            :return: The created table.
            """
            # Caption for the table
            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            runner = para.add_run(f"Table {str(table_count)}. Quality metrics across sequencing platforms.")
            runner.bold = True
            runner.italic = True
            
            # Create a table with one header row and 5 columns.
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            
            # Define header titles.
            headers = [
                'Sequencing Platform',
                'Genome Coverage (%)',
                'Ns in Sequence (%)',
                'Similarity (%)',
                'Read Coverage (Mean)'
            ]
            
            hdr_cells = table.rows[0].cells
            for i, title in enumerate(headers):
                hdr_cells[i].text = title
                # Center-align and bold the header text.
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add a row for each sequencing platform.
            for platform, metrics in aggregated_data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = platform
                row_cells[1].text = f"{metrics['coverage']:.1f}" if metrics['coverage'] is not None and metrics['similarity'] != 0 else "N/A"
                row_cells[2].text = f"{metrics['Ns']:.1f}" if metrics['Ns'] is not None and metrics['similarity'] != 0 else "N/A"
                row_cells[3].text = f"{metrics['similarity']:.1f}" if metrics['similarity'] is not None and metrics['similarity'] != 0 else "N/A"
                row_cells[4].text = f"{metrics['Mean coverage depth']:.1f}" if metrics['Mean coverage depth'] != "N/A" and metrics['Mean coverage depth']!=0 else "N/A"
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if " ".join(platform.split(" ")[:-1]) in user_platforms:
                    for i in range(5):
                        shading_elm = parse_xml(r'<w:shd {} w:fill="EFFA75"/>'.format(nsdecls('w')))
                        row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

            return table

        aggregated_data=aggregate_sample_data(sample_html_reports_original[sample], user_platform)
        create_platform_table(doc, aggregated_data, user_platform)

        last_paragraph = doc.add_paragraph() 
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        last_paragraph.add_run().add_break(WD_BREAK.PAGE)
        
        

       
























    # Appendixes
    appendix1=doc.add_paragraph()
    runner=appendix1.add_run(f"Appendix 1: Additional Information\n\n")
    runner.bold = True
    runner=appendix1.add_run(f"Samples provided and testing required\n")
    runner.bold = True
    runner.italic = True
    runner=appendix1.add_run(f"""Original samples distributed by UK NEQAS Microbiology were freeze dried with instructions on how to reconstitute the specimens.\n\n
Sequencing was carried out according to the laboratory's normal procedure.\n\n""")
    runner=appendix1.add_run(f"Data submission and analysis\n")
    runner.bold = True
    runner.italic = True
    runner=appendix1.add_run(f"""Data collection, quality control (QC), storage and analysis to WHO defined standards and requirements was carried out by UK NEQAS for Microbiology in collaborations with Cranfield University.\n\n""")
    runner=appendix1.add_run(f"""Participants were asked to submit FASTA, FASTQ, or BAM files for quality metric assessment. FASTQ files take priority over BAM files during analysis. If only a BAM file is provided, it  was realigned to the reference sequence. If no FASTA file was submitted, a consensus sequence was automatically generated from the FASTQ or BAM file. Otherwise, the submitted FASTA file was assumed to be the consensus sequence.\n\n""")
    runner=appendix1.add_run(f"Validated results\n")
    runner.bold = True
    runner.italic =   True
    runner=appendix1.add_run(f"""The EQA specimens were validated by a reference laboratory selected by UK NEQAS using Illumina sequencing, achieving 99.9% genome coverage at 20x depth.\nParticipant submissions were compared to the reference sequences, with lineages determined using Nextclade and Nextstrain identifiers as recommended by GISAID: EPI_ISL_412866 (RSV A) and EPI_ISL_1653999 (RSV B). This comparison covered all regions successfully sequenced by the reference laboratory, whose sequence quality was classified as very high.\n\n""")
    runner=appendix1.add_run(f"Quality Metrics\n")
    runner.bold = True
    runner.italic = True
    runner = appendix1.add_run(f"""Where appropriate data files have been provided by the participant, the following quality metrics have been stated on the reports:
➢  Genome Coverage (%) - The percentage of reference bases covered, with a threshold typically set at ≥90%. Computed using Nextclade.
➢  Ns in Sequence (%) - The percentage of ambiguous bases (N's) in the sequence, with a threshold typically set at ≤2%. Computed using Nextclade.
➢  Similarity (%) - The percentage of sequence similarity compared to the reference, with a threshold typically set at ≥95%. Computed using Nextclade.
➢  Read Coverage (Mean Depth) - The average depth of sequencing reads, with a threshold typically set at ≥50. Computed using Qualimap2.
NOTICE: A 100% similarity is not expected because Nextclade requires GISAID reference sequences for alignment and lineage assignment. \n\n\n""") 
    runner=appendix1.add_run(f"Participation and scoring submissions\n")
    runner.bold = True
    runner.italic = True

    appendix1.add_run("Enquiries: ").bold=True
    appendix1.runs[-1].font.size = Pt(9)
    appendix1.add_run("Pre-distribution test results are available should you experience a technical failure and wish to discuss "
        "the results. Written enquiries about this distribution should be addressed to Dr Sanjiv Rughooputh by email: ")
    appendix1.runs[-1].font.size = Pt(9)
    add_hyperlink(appendix1, "organiser@ukneqasmicro.org.uk", "organiser@ukneqasmicro.org.uk")
    appendix1.runs[-1].font.size = Pt(9)
    appendix1.add_run("\n\nAcknowledgements: ").bold=True 
    appendix1.runs[-1].font.size = Pt(9)
    appendix1.add_run("We would like to thank NICD, VIDRL and UKHSA for the provision of clinical isolates, UKHSA Manchester, "
        "VRD for their kind assistance with pre-distribution tests, and the Bioinformatics Group at Cranfield University's School of Engineering "
        "and Applied Sciences for bioinformatics analysis."
    )
    appendix1.runs[-1].font.size = Pt(9)
 
    # Save the document to a BytesIO object
    docx_io = BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return doc